import os
import json
from typing import Dict, Any, List
from datetime import datetime   
import mimetypes

# Pandas for Excel parsing
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# PDF parsing with multiple libraries for better compatibility
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import fitz  # PyMuPDF - better PDF text extraction
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Word parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel parsing with openpyxl (alternative to pandas)
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# CSV parsing
try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False

# JSON parsing (built-in)
import json as json_lib

# Create output directory for parsed documents
PARSED_DIR = "parsed_documents"
os.makedirs(PARSED_DIR, exist_ok=True)

async def parse_pdf(file_path: str) -> Dict[str, Any]:
    """Parse PDF file and extract text using multiple methods"""
    # Try PyMuPDF first (better text extraction)
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(file_path)
            pages = []
            full_text = ""
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                pages.append({
                    "page_number": page_num + 1,
                    "content": page_text.strip()
                })
                full_text += page_text + "\n"
            
            doc.close()
            
            # If no text was extracted, provide helpful message
            if not full_text.strip():
                full_text = f"""This PDF document ({os.path.basename(file_path)}) appears to contain scanned images or graphics rather than selectable text.

Content Details:
- Document has {total_pages} page(s)
- File size: {os.path.getsize(file_path) / 1024:.1f} KB
- No extractable text found

This could be because:
1. The document contains scanned images
2. The text is embedded as graphics
3. The document is password protected
4. The text encoding is not supported

To extract text from this document, you may need to:
- Use OCR (Optical Character Recognition) software
- Re-save the document with selectable text
- Convert images to text using specialized tools"""
            
            return {
                "type": "pdf",
                "total_pages": total_pages,
                "content": full_text.strip(),
                "pages": pages,
                "extraction_method": "PyMuPDF",
                "metadata": {
                    "file_name": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "character_count": len(full_text.strip()),
                    "has_extractable_text": len(full_text.strip()) > 0
                }
            }
        except Exception as e:
            print(f"⚠️ PyMuPDF failed: {str(e)}, trying PyPDF2...")

    # Fallback to PyPDF2 with similar enhancements
    if not PDF_AVAILABLE:
        return {"error": "PDF parsing not available. Install PyPDF2 or PyMuPDF: pip install PyPDF2 PyMuPDF"}
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            pages = []
            full_text = ""
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    pages.append({
                        "page_number": page_num + 1,
                        "content": page_text.strip()
                    })
                    full_text += page_text + "\n"
                except Exception as page_error:
                    print(f"⚠️ Error extracting page {page_num + 1}: {str(page_error)}")
                    pages.append({
                        "page_number": page_num + 1,
                        "content": f"Error extracting text from page {page_num + 1}",
                        "error": str(page_error)
                    })
            
            # Enhanced fallback message for PyPDF2
            if not full_text.strip():
                full_text = f"""PDF Analysis Report for {os.path.basename(file_path)}

Document Properties:
- Pages: {total_pages}
- File Size: {os.path.getsize(file_path) / 1024:.1f} KB
- Encrypted: {'Yes' if pdf_reader.is_encrypted else 'No'}

Text Extraction Status: No readable text found

This document may contain:
- Scanned images requiring OCR processing
- Graphics or charts without text layers
- Protected content
- Non-standard text encoding

The document structure has been preserved and can be used for further processing or manual review."""
            
            return {
                "type": "pdf",
                "total_pages": total_pages,
                "content": full_text.strip(),
                "pages": pages,
                "extraction_method": "PyPDF2",
                "metadata": {
                    "file_name": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "character_count": len(full_text.strip()),
                    "has_extractable_text": len(full_text.strip()) > 0,
                    "pdf_info": {
                        "encrypted": pdf_reader.is_encrypted,
                        "pages": total_pages
                    }
                }
            }
    
    except Exception as e:
        return {"error": f"PDF parsing failed: {str(e)}"}

async def parse_docx(file_path: str) -> Dict[str, Any]:
    """Parse Word document and extract text, tables, and formatting"""
    if not DOCX_AVAILABLE:
        return {"error": "Word parsing not available. Install python-docx: pip install python-docx"}
    
    try:
        doc = Document(file_path)
        paragraphs = []
        full_text = ""
        
        # Extract paragraphs with style information
        for para in doc.paragraphs:
            if para.text.strip():
                para_info = {
                    "text": para.text.strip(),
                    "style": para.style.name if para.style else "Normal"
                }
                paragraphs.append(para_info)
                full_text += para.text + "\n"
        
        # Extract tables with enhanced structure
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "table_number": table_idx + 1,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data["data"].append(row_data)
            
            tables.append(table_data)
        
        # Extract document properties
        props = doc.core_properties
        document_properties = {
            "title": props.title,
            "author": props.author,
            "subject": props.subject,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None
        }
        
        return {
            "type": "docx",
            "content": full_text.strip(),
            "paragraphs": paragraphs,
            "tables": tables,
            "document_properties": document_properties,
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "character_count": len(full_text.strip()),
                "word_count": len(full_text.split())
            }
        }
    
    except Exception as e:
        return {"error": f"Word document parsing failed: {str(e)}"}

async def parse_excel(file_path: str) -> Dict[str, Any]:
    """Parse Excel file with enhanced data extraction"""
    if PANDAS_AVAILABLE:
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            summary_stats = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Handle NaN values and convert to appropriate types
                df_cleaned = df.fillna("")
                
                # Get basic statistics
                numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
                stats = {}
                
                if numeric_columns:
                    for col in numeric_columns:
                        if not df[col].isna().all():
                            stats[col] = {
                                "mean": float(df[col].mean()) if not df[col].isna().all() else 0,
                                "min": float(df[col].min()) if not df[col].isna().all() else 0,
                                "max": float(df[col].max()) if not df[col].isna().all() else 0,
                                "count": int(df[col].count())
                            }
                
                sheets_data[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "data": df_cleaned.to_dict('records'),
                    "shape": {
                        "rows": len(df),
                        "columns": len(df.columns)
                    },
                    "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
                    "statistics": stats
                }
                
                summary_stats[sheet_name] = {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "numeric_columns": len(numeric_columns),
                    "empty_cells": int(df.isna().sum().sum())
                }
            
            return {
                "type": "excel",
                "sheets": sheets_data,
                "sheet_names": excel_file.sheet_names,
                "summary_statistics": summary_stats,
                "parsing_method": "pandas",
                "metadata": {
                    "file_name": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "sheet_count": len(excel_file.sheet_names),
                    "total_rows": sum(stats["rows"] for stats in summary_stats.values()),
                    "total_columns": sum(stats["columns"] for stats in summary_stats.values())
                }
            }
        except Exception as e:
            print(f"⚠️ Pandas Excel parsing failed: {str(e)}, trying openpyxl...")
    
    # Fallback to openpyxl
    if EXCEL_AVAILABLE:
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheets_data = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                data = []
                columns = []
                
                # Get all data from the sheet
                rows_data = list(sheet.iter_rows(values_only=True))
                
                if rows_data:
                    # Use first row as headers
                    columns = [str(cell) if cell is not None else f"Column_{i}" 
                              for i, cell in enumerate(rows_data[0])]
                    
                    # Process data rows
                    for row in rows_data[1:]:
                        row_data = {}
                        for i, value in enumerate(row):
                            if i < len(columns):
                                row_data[columns[i]] = str(value) if value is not None else ""
                        data.append(row_data)
                
                sheets_data[sheet_name] = {
                    "columns": columns,
                    "data": data,
                    "shape": {
                        "rows": len(data),
                        "columns": len(columns)
                    }
                }
            
            return {
                "type": "excel",
                "sheets": sheets_data,
                "sheet_names": list(workbook.sheetnames),
                "parsing_method": "openpyxl",
                "metadata": {
                    "file_name": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "sheet_count": len(workbook.sheetnames)
                }
            }
        except Exception as e:
            return {"error": f"Excel parsing with openpyxl failed: {str(e)}"}
    
    else:
        return {"error": "Excel parsing not available. Install pandas or openpyxl: pip install pandas openpyxl"}

async def parse_csv(file_path: str) -> Dict[str, Any]:
    """Parse CSV file with enhanced detection"""
    try:
        # Try to detect delimiter and encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read(10000)  # Read first 10KB
            
        # Try to detect encoding
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                content = raw_data.decode(encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if not content:
            return {"error": "Could not detect file encoding"}
        
        # Detect delimiter
        import csv
        sniffer = csv.Sniffer()
        delimiter = sniffer.sniff(content[:1000]).delimiter
        
        # Parse CSV
        data = []
        columns = []
        
        with open(file_path, 'r', encoding=used_encoding) as csvfile:
            csv_reader = csv.reader(csvfile, delimiter=delimiter)
            
            # Get headers
            columns = next(csv_reader, [])
            
            # Get data
            for row in csv_reader:
                if len(row) == len(columns):
                    row_data = {columns[i]: row[i] for i in range(len(columns))}
                    data.append(row_data)
        
        return {
            "type": "csv",
            "content": f"CSV file with {len(data)} rows and {len(columns)} columns",
            "columns": columns,
            "data": data,
            "delimiter": delimiter,
            "encoding": used_encoding,
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "row_count": len(data),
                "column_count": len(columns)
            }
        }
    except Exception as e:
        return {"error": f"CSV parsing failed: {str(e)}"}

async def parse_json(file_path: str) -> Dict[str, Any]:
    """Parse JSON file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            json_data = json_lib.load(file)
        
        # Analyze JSON structure
        def analyze_json_structure(obj, depth=0, max_depth=3):
            if depth > max_depth:
                return "..."
            
            if isinstance(obj, dict):
                return {k: analyze_json_structure(v, depth+1, max_depth) for k, v in list(obj.items())[:5]}
            elif isinstance(obj, list):
                return [analyze_json_structure(obj[0], depth+1, max_depth)] if obj else []
            else:
                return type(obj).__name__
        
        structure = analyze_json_structure(json_data);
        
        return {
            "type": "json",
            "content": json_lib.dumps(json_data, indent=2, ensure_ascii=False)[:2000] + "..." if len(str(json_data)) > 2000 else json_lib.dumps(json_data, indent=2, ensure_ascii=False),
            "data": json_data,
            "structure": structure,
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "json_type": type(json_data).__name__,
                "estimated_size": len(str(json_data))
            }
        }
    except json_lib.JSONDecodeError as e:
        return {"error": f"Invalid JSON format: {str(e)}"}
    except Exception as e:
        return {"error": f"JSON parsing failed: {str(e)}"}

async def parse_text(file_path: str) -> Dict[str, Any]:
    """Parse plain text file with encoding detection"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252', 'ascii']
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    used_encoding = encoding
                    break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return {"error": "Could not decode text file with any supported encoding"}
        
        lines = content.split('\n')
        
        # Basic text analysis
        word_count = len(content.split())
        sentence_count = content.count('.') + content.count('!') + content.count('?')
        paragraph_count = len([line for line in lines if line.strip()])
        
        return {
            "type": "text",
            "content": content,
            "lines": lines,
            "encoding": used_encoding,
            "analysis": {
                "word_count": word_count,
                "sentence_count": sentence_count,
                "paragraph_count": paragraph_count,
                "character_count": len(content),
                "character_count_no_spaces": len(content.replace(' ', ''))
            },
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "line_count": len(lines),
                "encoding_used": used_encoding
            }
        }
    
    except Exception as e:
        return {"error": f"Text file parsing failed: {str(e)}"}

async def run_document_parser_node(node_data: Dict[str, Any]) -> str:
    """Main function for enhanced document parser node"""
    try:
        file_path = node_data.get("file_path", "")
        
        if not file_path:
            return "Error: No file path provided"
        
        if not os.path.exists(file_path):
            return f"Error: File not found at path: {file_path}"
        
        # Detect file type
        mime_type, _ = mimetypes.guess_type(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        print(f"📄 Parsing document: {os.path.basename(file_path)}")
        print(f"📋 File type: {file_ext}, MIME: {mime_type}")
        print(f"📊 File size: {os.path.getsize(file_path) / 1024:.1f} KB")
        
        # Route to appropriate parser
        if file_ext == '.pdf':
            result = await parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            if file_ext == '.doc':
                return "Error: .doc files not supported. Please convert to .docx format"
            result = await parse_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            if file_ext == '.xls':
                return "Error: .xls files not supported. Please convert to .xlsx format"
            result = await parse_excel(file_path)
        elif file_ext == '.csv':
            result = await parse_csv(file_path)
        elif file_ext == '.json':
            result = await parse_json(file_path)
        elif file_ext in ['.txt', '.md', '.log']:
            result = await parse_text(file_path)
        else:
            return f"Error: Unsupported file type: {file_ext}. Supported types: PDF, DOCX, XLSX, CSV, JSON, TXT, MD"
        
        if "error" in result:
            print(f"❌ Parse error: {result['error']}")
            return result["error"]
        
        # Save parsed data to JSON file for downstream nodes
        base_filename = os.path.splitext(os.path.basename(file_path))[0]
        output_filename = f"parsed_{base_filename}_{int(datetime.now().timestamp())}.json"
        output_path = os.path.join(PARSED_DIR, output_filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False, default=str)
        
        # Create summary
        doc_type = result.get("type", "unknown")
        metadata = result.get("metadata", {})
        
        print(f"✅ Document parsed successfully!")
        print(f"📁 Type: {doc_type.upper()}")
        print(f"📏 Content length: {len(result.get('content', ''))} characters")
        print(f"💾 JSON saved to: {output_path}")
        
        # Enhanced summary based on document type
        if doc_type == "pdf":
            print(f"📖 Pages: {result.get('total_pages', 0)}")
        elif doc_type == "excel":
            print(f"📊 Sheets: {len(result.get('sheet_names', []))}")
        elif doc_type == "docx":
            print(f"📝 Paragraphs: {metadata.get('paragraph_count', 0)}, Tables: {metadata.get('table_count', 0)}")
        
        return f"Document parsed: {output_path}"
        
    except Exception as e:
        error_msg = f"Document parsing failed: {str(e)}"
        print(f"❌ {error_msg}")
        return error_msg

# Helper function to get parsing capabilities
def get_parsing_capabilities():
    """Return available parsing capabilities"""
    capabilities = {
        "pdf": {
            "available": PDF_AVAILABLE or PYMUPDF_AVAILABLE,
            "libraries": []
        },
        "docx": {
            "available": DOCX_AVAILABLE,
            "libraries": ["python-docx"] if DOCX_AVAILABLE else []
        },
        "excel": {
            "available": PANDAS_AVAILABLE or EXCEL_AVAILABLE,
            "libraries": []
        },
        "csv": {
            "available": True,
            "libraries": ["built-in"]
        },
        "json": {
            "available": True,
            "libraries": ["built-in"]
        },
        "text": {
            "available": True,
            "libraries": ["built-in"]
        }
    }
    
    if PDF_AVAILABLE:
        capabilities["pdf"]["libraries"].append("PyPDF2")
    if PYMUPDF_AVAILABLE:
        capabilities["pdf"]["libraries"].append("PyMuPDF")
    
    if PANDAS_AVAILABLE:
        capabilities["excel"]["libraries"].append("pandas")
    if EXCEL_AVAILABLE:
        capabilities["excel"]["libraries"].append("openpyxl")
    
    return capabilities

# Test function
async def test_document_parser():
    """Test document parser with sample files"""
    print("🧪 Testing document parser capabilities...")
    capabilities = get_parsing_capabilities()
    
    for doc_type, info in capabilities.items():
        status = "✅" if info["available"] else "❌"
        libraries = ", ".join(info["libraries"]) if info["libraries"] else "None"
        print(f"{status} {doc_type.upper()}: {libraries}")
