import os
import uuid
import json
from datetime import datetime
from typing import Dict, Any, Optional, List
import re

# PDF generation
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib import colors
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie
    from reportlab.graphics.charts.linecharts import HorizontalLineChart
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

OUTPUT_DIR = "generated_reports"
# Ensure directory exists with absolute path
if not os.path.isabs(OUTPUT_DIR):
    OUTPUT_DIR = os.path.join(os.getcwd(), OUTPUT_DIR)
os.makedirs(OUTPUT_DIR, exist_ok=True)

class ReportStyles:
    """Define custom styles for reports"""
    
    @staticmethod
    def get_pdf_styles():
        """Get enhanced PDF styles"""
        styles = getSampleStyleSheet()
        
        # Enhanced title style
        title_style = ParagraphStyle(
            'EnhancedTitle',
            parent=styles['Title'],
            fontSize=28,
            spaceAfter=30,
            spaceBefore=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1f4e79'),
            fontName='Helvetica-Bold'
        )
        
        # Subtitle style
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=16,
            spaceAfter=20,
            spaceBefore=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#2e75b6'),
            fontName='Helvetica-Oblique'
        )
        
        # Enhanced heading styles
        h1_style = ParagraphStyle(
            'EnhancedHeading1',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=15,
            spaceBefore=25,
            textColor=colors.HexColor('#1f4e79'),
            borderWidth=1,
            borderColor=colors.HexColor('#2e75b6'),
            borderPadding=5,
            backColor=colors.HexColor('#f2f8ff')
        )
        
        h2_style = ParagraphStyle(
            'EnhancedHeading2',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.HexColor('#2e75b6'),
            borderWidth=0,
            leftIndent=10
        )
        
        h3_style = ParagraphStyle(
            'EnhancedHeading3',
            parent=styles['Heading3'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=15,
            textColor=colors.HexColor('#4472c4'),
            leftIndent=20
        )
        
        # Enhanced body text
        body_style = ParagraphStyle(
            'EnhancedBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=8,
            spaceBefore=4,
            alignment=TA_JUSTIFY,
            textColor=colors.HexColor('#333333')
        )
        
        # Quote/highlight style
        quote_style = ParagraphStyle(
            'Quote',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            leftIndent=30,
            rightIndent=30,
            spaceAfter=12,
            spaceBefore=12,
            backColor=colors.HexColor('#f8f9fa'),
            borderWidth=1,
            borderColor=colors.HexColor('#dee2e6'),
            borderPadding=10,
            fontName='Helvetica-Oblique'
        )
        
        # Code/data style
        code_style = ParagraphStyle(
            'Code',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Courier',
            backColor=colors.HexColor('#f8f9fa'),
            borderWidth=1,
            borderColor=colors.HexColor('#e9ecef'),
            borderPadding=8,
            leftIndent=20
        )
        
        return {
            'title': title_style,
            'subtitle': subtitle_style,
            'h1': h1_style,
            'h2': h2_style,
            'h3': h3_style,
            'body': body_style,
            'quote': quote_style,
            'code': code_style
        }

def create_enhanced_table(data: List[List], headers: List[str] = None, table_style: str = "default"):
    """Create an enhanced table with styling"""
    if not data:
        return None
    
    # Prepare table data
    if headers:
        table_data = [headers] + data
    else:
        table_data = data
    
    # Calculate column widths based on content
    max_cols = max(len(row) for row in table_data)
    col_widths = [2 * inch] * max_cols
    
    table = Table(table_data, colWidths=col_widths)
    
    # Apply styling based on table_style
    if table_style == "financial":
        style_commands = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4e79')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
        ]
    elif table_style == "data":
        style_commands = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e75b6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]
    else:  # default
        style_commands = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472c4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6'))
        ]
    
    table.setStyle(TableStyle(style_commands))
    return table

def create_summary_box(title: str, content: str, color_scheme: str = "blue"):
    """Create a styled summary box"""
    color_schemes = {
        "blue": {"bg": colors.HexColor('#e3f2fd'), "border": colors.HexColor('#1976d2')},
        "green": {"bg": colors.HexColor('#e8f5e8'), "border": colors.HexColor('#4caf50')},
        "orange": {"bg": colors.HexColor('#fff3e0'), "border": colors.HexColor('#ff9800')},
        "red": {"bg": colors.HexColor('#ffebee'), "border": colors.HexColor('#f44336')}
    }
    
    scheme = color_schemes.get(color_scheme, color_schemes["blue"])
    
    box_style = ParagraphStyle(
        'SummaryBox',
        fontSize=10,
        leading=14,
        spaceAfter=15,
        spaceBefore=15,
        backColor=scheme["bg"],
        borderWidth=2,
        borderColor=scheme["border"],
        borderPadding=15,
        leftIndent=10,
        rightIndent=10
    )
    
    title_style = ParagraphStyle(
        'SummaryBoxTitle',
        fontSize=12,
        fontName='Helvetica-Bold',
        textColor=scheme["border"],
        spaceAfter=8
    )
    
    return [
        Paragraph(title, title_style),
        Paragraph(content, box_style)
    ]

def parse_enhanced_content(content: str, styles: Dict) -> List:
    """Parse content with enhanced formatting"""
    story = []
    lines = content.split('\n')
    import re

    def safe_bold(text):
        # Replace **bold** with <b>bold</b> safely
        return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            story.append(Spacer(1, 8))
            i += 1
            continue
        
        # Enhanced headers
        if line.startswith('# '):
            header_text = line[2:].strip()
            story.append(Paragraph(header_text, styles['h1']))
        elif line.startswith('## '):
            header_text = line[3:].strip()
            story.append(Paragraph(header_text, styles['h2']))
        elif line.startswith('### '):
            header_text = line[4:].strip()
            story.append(Paragraph(header_text, styles['h3']))
        
        # Bullet points with sub-bullets
        elif line.startswith('- '):
            bullet_text = line[2:].strip()
            bullet_style = ParagraphStyle(
                'Bullet',
                parent=styles['body'],
                leftIndent=20,
                bulletIndent=10,
                bulletFontName='Symbol',
                bulletText='•'
            )
            story.append(Paragraph(bullet_text, bullet_style))
        
        elif line.startswith('  - '):
            sub_bullet_text = line[4:].strip()
            sub_bullet_style = ParagraphStyle(
                'SubBullet',
                parent=styles['body'],
                leftIndent=40,
                bulletIndent=30,
                bulletFontName='Symbol',
                bulletText='◦'
            )
            story.append(Paragraph(sub_bullet_text, sub_bullet_style))
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            numbered_text = re.sub(r'^\d+\.\s', '', line)
            numbered_style = ParagraphStyle(
                'Numbered',
                parent=styles['body'],
                leftIndent=20,
                bulletIndent=10
            )
            story.append(Paragraph(numbered_text, numbered_style))
        
        # Code blocks
        elif line.startswith('```'):
            # Multi-line code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_content = '\n'.join(code_lines)
                story.append(Paragraph(f"<font name='Courier'>{code_content}</font>", styles['code']))
        
        # Quotes
        elif line.startswith('> '):
            quote_text = line[2:].strip()
            story.append(Paragraph(quote_text, styles['quote']))
        
        # Bold formatting (only **bold**)
        elif '**' in line:
            formatted_line = safe_bold(line)
            story.append(Paragraph(formatted_line, styles['body']))
        
        # Regular paragraph
        else:
            story.append(Paragraph(line, styles['body']))
        
        i += 1

    return story

def generate_pdf_report(title: str, content: str, data: Optional[Dict] = None) -> str:
    """Generate a professional PDF report using ReportLab"""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("ReportLab not available. Install with: pip install reportlab")
    
    try:
        filename = f"report_{uuid.uuid4().hex[:8]}.pdf"
        file_path = os.path.join(OUTPUT_DIR, filename)
        
        # Create document with custom page template
        doc = SimpleDocTemplate(
            file_path, 
            pagesize=A4,
            topMargin=1*inch,
            bottomMargin=1*inch,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch
        )
        
        story = []
        styles = ReportStyles.get_pdf_styles()
        
        # Title page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(title, styles['title']))
        story.append(Spacer(1, 0.5*inch))
        
        # Subtitle with generation info
        subtitle_text = f"Generated by AutoFlow on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        story.append(Paragraph(subtitle_text, styles['subtitle']))
        story.append(Spacer(1, 1*inch))
        
        # Executive Summary Box (if data available)
        if data and any(key.startswith('ai_response') for key in data.keys()):
            summary_content = []
            for key, value in data.items():
                if key.startswith('ai_response') and isinstance(value, str):
                    summary_content.append(value[:200] + "..." if len(value) > 200 else value)
            
            if summary_content:
                exec_summary = " ".join(summary_content)
                story.extend(create_summary_box("Executive Summary", exec_summary, "blue"))
        
        story.append(PageBreak())
        
        # Table of Contents (simplified)
        story.append(Paragraph("Table of Contents", styles['h1']))
        toc_style = ParagraphStyle('TOC', parent=styles['body'], leftIndent=20)
        
        # Extract headers from content for TOC
        headers = []
        for line in content.split('\n'):
            if line.startswith('# '):
                headers.append(("1. " + line[2:].strip(), 1))
            elif line.startswith('## '):
                headers.append(("   1.1 " + line[3:].strip(), 2))
        
        for header_text, level in headers:
            story.append(Paragraph(header_text, toc_style))
        
        story.append(PageBreak())
        
        # Main content with enhanced formatting
        story.extend(parse_enhanced_content(content, styles))
        
        # Data sections
        if data and isinstance(data, dict):
            story.append(PageBreak())
            story.append(Paragraph("Data Analysis & Metrics", styles['h1']))
            
            # Workflow metrics
            workflow_metrics = [
                ["Metric", "Value"],
                ["Execution Time", data.get("workflow_execution_time", "N/A")],
                ["Nodes Processed", str(data.get("total_nodes_processed", 0))],
                ["Report Format", "PDF"],
                ["Generated By", "AutoFlow Report Generator"]
            ]
            
            metrics_table = create_enhanced_table(workflow_metrics[1:], workflow_metrics[0], "financial")
            if metrics_table:
                story.append(Spacer(1, 20))
                story.append(Paragraph("Workflow Metrics", styles['h2']))
                story.append(metrics_table)
                story.append(Spacer(1, 20))
            
            # File analysis data
            file_data = []
            for key, value in data.items():
                if key.startswith(('document_', 'uploaded_file_')):
                    if isinstance(value, dict):
                        file_info = [
                            ["Property", "Value"],
                            ["File Name", key.replace('document_', '').replace('uploaded_file_', '')],
                            ["Type", value.get('type', 'Unknown')],
                            ["Size", value.get('size', 'Unknown')]
                        ]
                        if value.get('url'):
                            file_info.append(["URL", value.get('url', '')[:50] + "..."])
                        
                        file_table = create_enhanced_table(file_info[1:], file_info[0], "data")
                        if file_table:
                            story.append(Paragraph(f"File Analysis: {key}", styles['h2']))
                            story.append(file_table)
                            story.append(Spacer(1, 15))
        
        # Footer
        story.append(PageBreak())
        story.append(Spacer(1, 2*inch))
        
        footer_content = """
        This report was automatically generated by AutoFlow, an intelligent workflow automation platform. 
        The data presented in this report has been processed and analyzed using various AI models and document 
        parsing technologies to provide comprehensive insights and summaries.
        """
        
        footer_style = ParagraphStyle(
            'Footer',
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#666666'),
            spaceAfter=20
        )
        
        story.append(Paragraph(footer_content, footer_style))
        story.append(Paragraph("— End of Report —", footer_style))
        
        # Build the document
        doc.build(story)
        return file_path
        
    except Exception as e:
        raise Exception(f"Enhanced PDF generation failed: {str(e)}")

def generate_docx_report(title: str, content: str, data: Optional[Dict] = None) -> str:
    """Generate an enhanced Word document report"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available. Install with: pip install python-docx")
    
    try:
        filename = f"report_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(OUTPUT_DIR, filename)
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title page
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(31, 78, 121)
        
        # Subtitle
        subtitle_text = f"Generated by AutoFlow on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        subtitle_para = doc.add_paragraph(subtitle_text)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(46, 117, 182)
        
        doc.add_page_break()
        
        # Executive Summary
        if data and any(key.startswith('ai_response') for key in data.keys()):
            doc.add_heading('Executive Summary', level=1)
            summary_content = []
            for key, value in data.items():
                if key.startswith('ai_response') and isinstance(value, str):
                    summary_content.append(value[:300] + "..." if len(value) > 300 else value)
            
            if summary_content:
                exec_summary = " ".join(summary_content)
                summary_para = doc.add_paragraph(exec_summary)
                summary_para.style = 'Intense Quote'
        
        doc.add_page_break()
        
        # Enhanced content parsing
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('  - '):
                para = doc.add_paragraph(line[4:], style='List Bullet 2')
            elif re.match(r'^\d+\.\s', line):
                numbered_text = re.sub(r'^\d+\.\s', '', line)
                para = doc.add_paragraph(numbered_text, style='List Number')
            elif line.startswith('> '):
                para = doc.add_paragraph(line[2:], style='Intense Quote')
            else:
                # Handle formatting
                para = doc.add_paragraph()
                if '**' in line:
                    parts = line.split('**')
                    for j, part in enumerate(parts):
                        run = para.add_run(part)
                        if j % 2 == 1:  # Odd indices are bold
                            run.bold = True
                else:
                    para.add_run(line)
            i += 1
        
        doc.add_page_break()
        
        # Data analysis section
        if data and isinstance(data, dict):
            doc.add_heading('Data Analysis & Metrics', level=1)
            
            # Workflow metrics table
            doc.add_heading('Workflow Metrics', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            metrics = [
                ("Execution Time", data.get("workflow_execution_time", "N/A")),
                ("Nodes Processed", str(data.get("total_nodes_processed", 0))),
                ("Report Format", "DOCX"),
                ("Generated By", "AutoFlow Report Generator")
            ]
            
            for metric, value in metrics:
                row_cells = table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
            
            # File analysis tables
            for key, value in data.items():
                if key.startswith(('document_', 'uploaded_file_')) and isinstance(value, dict):
                    doc.add_heading(f'File Analysis: {key}', level=3)
                    
                    file_table = doc.add_table(rows=1, cols=2)
                    file_table.style = 'Table Grid'
                    
                    hdr_cells = file_table.rows[0].cells
                    hdr_cells[0].text = 'Property'
                    hdr_cells[1].text = 'Value'
                    
                    file_props = [
                        ("Type", value.get('type', 'Unknown')),
                        ("Size", value.get('size', 'Unknown'))
                    ]
                    
                    if value.get('url'):
                        file_props.append(("URL", value.get('url', '')[:50] + "..."))
                    
                    for prop, val in file_props:
                        row_cells = file_table.add_row().cells
                        row_cells[0].text = prop
                        row_cells[1].text = str(val)
        
        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            "This report was automatically generated by AutoFlow, an intelligent workflow automation platform."
        )
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(102, 102, 102)
        
        doc.save(file_path)
        return file_path
        
    except Exception as e:
        raise Exception(f"Enhanced DOCX generation failed: {str(e)}")

async def run_report_generator_node(node_data: Dict[str, Any]) -> str:
    """Main function for enhanced report generator node"""
    try:
        title = node_data.get("title", "AutoFlow Comprehensive Report")
        content = node_data.get("content", "# AutoFlow Report\n\nThis comprehensive report was generated automatically.")
        format_type = node_data.get("format", "pdf").lower()
        
        print(f"📊 Generating enhanced {format_type.upper()} report: {title}")
        print(f"📄 Content length: {len(content)} characters")
        
        # Extract and enhance report data
        report_data = node_data.get("data", {})
        
        # Add metadata to report data
        report_data.update({
            "report_generation_time": datetime.now().isoformat(),
            "report_title": title,
            "report_format": format_type.upper(),
            "content_sections": len([line for line in content.split('\n') if line.startswith('#')])
        })
        
        if format_type == "pdf":
            file_path = generate_pdf_report(title, content, report_data)
        elif format_type == "docx":
            file_path = generate_docx_report(title, content, report_data)
        else:
            return f"Error: Unsupported format '{format_type}'. Use 'pdf' or 'docx'"
        
        file_size = os.path.getsize(file_path) / 1024  # Size in KB
        print(f"✅ Enhanced report generated successfully!")
        print(f"📁 File: {os.path.basename(file_path)} ({file_size:.1f} KB)")
        print(f"📂 Location: {file_path}")
        
        return f"Report generated: {file_path}"
        
    except ImportError as e:
        return f"Error: {str(e)}"
    except Exception as e:
        print(f"❌ Report generation error: {str(e)}")
        return f"Report generation failed: {str(e)}"
